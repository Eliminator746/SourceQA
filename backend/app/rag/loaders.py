import os
import tempfile
from io import BytesIO

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader


def load_pdf(
    contents: bytes,
    document_id: str,
    user_id: str,
    filename: str
) -> list[Document]:

    temp_path = os.path.join(
        tempfile.gettempdir(),
        f"{document_id}.pdf"
    )

    try:
        with open(temp_path, "wb") as file:
            file.write(contents)

        loader = PyPDFLoader(temp_path)

        documents = loader.load() # creates one Document per page

        for document in documents:
            document.metadata.update({
                "document_id": document_id,
                "user_id": user_id,
                "filename": filename,
                "file_type": "pdf"
            })

        return documents

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def load_docx(
    contents: bytes,
    document_id: str,
    user_id: str,
    filename: str
) -> list[Document]:

    docx_file = BytesIO(contents)

    doc = DocxDocument(docx_file)

    documents = []

    for index, paragraph in enumerate(doc.paragraphs):

        text = paragraph.text.strip()

        if not text:
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "document_id": document_id,
                    "user_id": user_id,
                    "filename": filename,
                    "file_type": "docx",
                    "paragraph": index
                }
            )
        )

    return documents


def load_txt(
    contents: bytes,
    document_id: str,
    user_id: str,
    filename: str
) -> list[Document]:

    text = contents.decode(
        "utf-8",
        errors="ignore"
    )

    return [
        Document(
            page_content=text,
            metadata={
                "document_id": document_id,
                "user_id": user_id,
                "filename": filename,
                "file_type": "txt"
            }
        )
    ]


def load_document(
    contents: bytes,
    file_type: str,
    document_id: str,
    user_id: str,
    filename: str
) -> list[Document]:

    if file_type == "pdf":
        return load_pdf(
            contents,
            document_id,
            user_id,
            filename
        )

    if file_type == "docx":
        return load_docx(
            contents,
            document_id,
            user_id,
            filename
        )

    if file_type == "txt":
        return load_txt(
            contents,
            document_id,
            user_id,
            filename
        )

    raise ValueError(
        f"Unsupported file type: {file_type}"
    )
    
    
# S3 / Upload bytes
#        ↓
# temporary PDF
#        ↓
# PyPDFLoader
#        ↓
# LangChain Documents
#        ↓
# temporary PDF deleted